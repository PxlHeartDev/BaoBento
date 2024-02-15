# Library imports
from tkinter import *
import time
import json
from docx import Document
from docx.shared import Inches
import os
import math

# Explicit imports
from tkinter import messagebox, ttk
from tkinter.simpledialog import askinteger, askstring
from time import strftime

# Me imports
from main import db, cursor, getCustomerData, refreshCustomerData
from gui import root, createText, createButton, createEntryBox, createDropdown, createCheckbox
from sql import retrieveCustomerOrders, deleteCustomer, updateDetails, updateEmail, updateNotif, updatePassword
from orders import keyFromVal, keyFromValPrecise, indexFromVal, Appetiser, Bao, Bento, Classic, Side


# Global, for use with marking orders
global selectedOrder
selectedOrder = []

# Customer settings menu
def createCustomerSettingsToplevel():
    # Notification prefences variables
    emailNotif = IntVar()
    SMSNotif = IntVar()

    # Get the data of the logged in customer
    data = getCustomerData()
    
    # Get existing notification preferences
    notifPref = bin(data[6]).replace("0b", "").zfill(2)
    emailNotif.set(notifPref[1])
    SMSNotif.set(notifPref[0])

    # Create top level
    CustomerSettings = Toplevel(root, bg='#000000')
    CustomerSettings.geometry('1060x690')
    createText([CustomerSettings], 1, 0, 4, "Settings", "Calibri 35 bold")

    # Notification preferences
    createText([CustomerSettings], 2, 0, 1, "Notification Preferences: ", pady=5, padx=38)
    createText([CustomerSettings], 3, 0, 1, "Note: This preference only affects promotional messages\nConfirmation messages regarding orders or otherwise will always be sent", "Calibri 12", pady=5, padx=38, sticky='n')
    e = createCheckbox([CustomerSettings], 2, 1, 1, "Email", emailNotif, initVal=(notifPref[1] == "1"))
    s = createCheckbox([CustomerSettings], 2, 2, 1, "SMS", SMSNotif, initVal=(notifPref[0] == "1"))
    createButton([CustomerSettings], 3, 1, 1, "Confirm Preferences", lambda:updateNotif(data[0], SMSNotif.get(), emailNotif.get(), CustomerSettings), "Calibri 15", width=18, padx=5, pady=8, sticky='n')
    

    ### Edit details section ###

    # Re-use entry boxes
    from gui import firstName, lastName, number, email, password, conPassword

    # Set the entry boxes
    firstName.set(data[1])
    lastName.set(data[2])
    number.set(data[3])
    email.set(data[4])

    # Personal details
    createText([CustomerSettings], 4, 0, 1, "First Name: ", pady=5, padx=38)
    createEntryBox([CustomerSettings], 4, 1, 1, firstName, width=21)
    createText([CustomerSettings], 5, 0, 1, "Last Name: ", pady=5, padx=38)
    createEntryBox([CustomerSettings], 5, 1, 1, lastName, width=21)
    createText([CustomerSettings], 6, 0, 1, "Ph. Number: ", pady=5, padx=38)
    createEntryBox([CustomerSettings], 6, 1, 1, number, width=21)
    createButton([CustomerSettings], 5, 2, 1, "Update Details", lambda:updateDetails(data[0], firstName.get(), lastName.get(), number.get(), CustomerSettings), "Calibri 15", width=15, padx=5)

    # Email
    createText([CustomerSettings], 7, 0, 1, "Email: ", pady=20, padx=38)
    createEntryBox([CustomerSettings], 7, 1, 1, email, font="Calibri 15", width=31, ipady=9)
    createButton([CustomerSettings], 7, 2, 1, "Update Email", lambda:updateEmail(data[0], email.get(), CustomerSettings), "Calibri 15", width=15, padx=5)

    # Password
    createText([CustomerSettings], 8, 0, 1, "New Password: ", pady=5, padx=38)
    createEntryBox([CustomerSettings], 8, 1, 1, password, width=21, password=True)
    createText([CustomerSettings], 9, 0, 1, "Old Password: ", pady=5, padx=38)
    createEntryBox([CustomerSettings], 9, 1, 1, conPassword, width=21, password=True)
    createButton([CustomerSettings], 8, 2, 1, "Update Password", lambda:updatePassword(data[0], password.get(), conPassword.get(), CustomerSettings), "Calibri 15", width=15, padx=5, rowspan=2)

    # Delete account
    createButton([CustomerSettings], 10, 0, 4, "Delete Account", lambda:deleteCustomer(getCustomerData()[0], CustomerSettings), "Calibri 18", width=15)
    
    # Back
    createButton([CustomerSettings], 11, 0, 4, "Back", lambda:CustomerSettings.destroy(), "Calibri 18", pady=5, width=20)

#############################################################################################################################
#############################################################################################################################
#############################################################################################################################

# Owner menu for creating a new order
def createOwnerCreateOrderTopLevel():
    # Create top level
    OwnerCreateOrder = Toplevel(root, bg='#000000')
    OwnerCreateOrder.geometry('960x690')

    # Title and buttons
    createText([OwnerCreateOrder], 1, 0, 4, "Create Order", "Calibri 35 bold")
    createButton([OwnerCreateOrder], 2, 0, 4, "Appetisers", lambda:createOwnerAddItemTopLevel("appetisers"), "Calibri 30")
    createButton([OwnerCreateOrder], 3, 0, 4, "Baos", lambda:createOwnerAddItemTopLevel("baos"), "Calibri 30")
    createButton([OwnerCreateOrder], 4, 0, 4, "Bentos", lambda:createOwnerAddItemTopLevel("bentos"), "Calibri 30")
    createButton([OwnerCreateOrder], 5, 0, 4, "Classics", lambda:createOwnerAddItemTopLevel("classics"), "Calibri 30")
    createButton([OwnerCreateOrder], 6, 0, 4, "Sides", lambda:createOwnerAddItemTopLevel("sides"), "Calibri 30")


    def completeOrder():
        o.completeOrder(
            askinteger("Pickup time", "Please enter a pickup time in seconds since Jan 1 2024, 00:00 UTC+0"),
            askstring("Notes", "Please enter any notes about the order, such as allergens or specific requests")
        )

    createButton([OwnerCreateOrder], 7, 0, 4, "View order", lambda:createOrderReceipt(OwnerCreateOrder, o.getOrderContent()))
    createButton([OwnerCreateOrder], 8, 0, 4, "Complete order", completeOrder)


    # Import here for performance and to circumvent circular import issues
    from orders import Order, modTypes, sauceDict

    # Order object
    o = Order()

    # Sub-function, creates the add item top level
    def createOwnerAddItemTopLevel(type: str):
        # Create top level
        OwnerAddItem = Toplevel(root, bg='#000000')
        OwnerAddItem.geometry('1200x960')

        # Check the item type
        match(type):
            case("appetisers"):

                # Import here because idk honestly man I'm tired, probably circular import circumvention or something
                from orders import appetisers, appetiserPermittedSauces

                # Store the auto-generated UI elements and variables in a list so they can be easily accessed
                UIElements = []
                vars = []
                
                # Sub-sub-function, runs when the appetiser in the dropdown is changed
                def appetiserChanged(*args):
                    # Ignore if it is the init value
                    if('selected' in appetiserSelected.get()): return
                    # Get the appetiser details in full
                    appStuff = appetisers[keyFromVal(appetisers, appetiserSelected.get())]
                    # Re-used options list
                    oListMod = []
                    
                    # Destroy each auto-generated UI element
                    [e.destroy() for e in UIElements]
                    
                    # Clear the lists
                    UIElements.clear()
                    vars.clear()

                    for i in range(0, len(appStuff['mod'])):
                        m = appStuff['mod'][i]
                        oListMod.clear()
                        vars.append(StringVar())
                        for v in modTypes[m['modType']].values():
                            oListMod.append(v)
                        UIElements.append(createText([OwnerAddItem], 5+i, 0, 1, f"{m['name']}: "))
                        UIElements.append(createDropdown([OwnerAddItem], 5+i, 1, 1, oListMod, vars[i], "Calibri 15 bold", modTypes[m['modType']][m['default']], 24))

                    if(appStuff['defaultSauce'] != -1):
                        oListSauce = []
                        for i in appetiserPermittedSauces:
                            oListSauce.append(sauceDict[i])
                        if(appStuff['defaultSauce'] == 3):
                            oListSauce[2] = "Small Curry"
                        UIElements.append(createDropdown([OwnerAddItem], 5, 2, 1, oListSauce, sauceSelected, "Calibri 15 bold", sauceDict[appStuff['defaultSauce']], 24, 5))
                    else:
                        sauceSelected.set("")
                        UIElements.append(createText([OwnerAddItem], 5, 2, 1, "No sauce", "Calibri 20"))
                    
                    UIElements.append(createButton([OwnerAddItem], 9, 1, 4, "Add Item", addAppetiser, "Calibri 20", pady=5))

                def addAppetiser():
                    o.addItem(Appetiser(
                        count.get(),
                        keyFromVal(appetisers, appetiserSelected.get()),
                        [keyFromVal(modTypes[keyFromVal(modTypes, m.get())], m.get()) for m in vars],
                        keyFromVal(sauceDict, sauceSelected.get()),
                        note.get()))
                    messagebox.showinfo("Success", "Successfully Added appetiser to order", parent=OwnerAddItem)
                    OwnerAddItem.destroy()

                createText([OwnerAddItem], 1, 0, 4, "Appetisers", "Calibri 35 bold")
                oListApp = []
                for a in appetisers.values():
                    oListApp.append(a['name'])
                appetiserSelected = StringVar()
                sauceSelected = StringVar()
                appetiserSelected.trace_add("write", appetiserChanged)
                createText([OwnerAddItem], 2, 0, 1, "Appetiser: ")
                createDropdown([OwnerAddItem], 2, 1, 1, oListApp, appetiserSelected, "Calibri 15 bold", "No appetiser selected", 24)

                count = IntVar()
                createDropdown([OwnerAddItem], 2, 2, 1, [1, 2, 3, 4, 5], count, "Calibri 15 bold", 1, 5)

                createText([OwnerAddItem], 4, 0, 2, "Modifiers: ", sticky='s', pady=10)
                createText([OwnerAddItem], 4, 2, 2, "Sauce: ", sticky='s', pady=10)

                note = StringVar()
                createText([OwnerAddItem], 8, 0, 1, "Notes: ", "Calibri 20", pady=10)
                createEntryBox([OwnerAddItem], 8, 1, 2, note, "Calibri 20", width=30, ipadx=30)

            case("baos"):
                from orders import baos, baoPermittedSauces, picklesDict

                UIElements = []
                sauceAm = []
                vars = []

                def sauceChanged(*args):
                    [s.destroy() for s in sauceAm]
                    sauceAm.clear()
                    if('No' in sauceSelected.get()):
                        return
                    sauceAm.append(createDropdown([OwnerAddItem], 6, 2, 1, ["Less", "With", "Extra"], vars[0], "Calibri 15 bold", 'With', 15))

                def baoChanged(*args):
                    if('selected' in baoSelected.get()):
                        return
                    baoStuff = baos[keyFromVal(baos, baoSelected.get())]
                    
                    [e.destroy() for e in UIElements]
                        
                    UIElements.clear()
                    vars.clear()

                    vars.append(StringVar())
                        
                    oListPickles = []
                    for v in modTypes[1].values():
                        oListPickles.append(v)

                    for i in range(0, 5):
                        vars.append(StringVar())
                        UIElements.append(createText([OwnerAddItem], 5+i, 0, 1, f"{picklesDict[i+1]}: "))
                        UIElements.append(createDropdown([OwnerAddItem], 5+i, 1, 1, oListPickles, vars[i+1], "Calibri 15 bold", modTypes[1][baoStuff['pickles'][i]], 8))

                    oListSauce = []
                    for i in baoPermittedSauces:
                        oListSauce.append(sauceDict[i])
                    UIElements.append(createDropdown([OwnerAddItem], 5, 2, 1, oListSauce, sauceSelected, "Calibri 15 bold", sauceDict[baoStuff['sauce']], 15, 5))
                    
                    UIElements.append(createButton([OwnerAddItem], 12, 1, 4, "Add Item", addBao, "Calibri 20", pady=5))


                def addBao():
                    o.addItem(Bao
                            (1, 
                            keyFromVal(baos, baoSelected.get()),
                            keyFromVal(sauceDict, sauceSelected.get()),
                            keyFromVal(modTypes[1], vars[0].get()),
                            [keyFromVal(modTypes[1], m.get()) for m in vars[1:]],
                            note.get()))
                    messagebox.showinfo("Success", "Successfully added Bao to order", parent=OwnerAddItem)
                    OwnerAddItem.destroy()

                def noPickles():
                    for v in vars[1:]:
                        v.set("No")

                createText([OwnerAddItem], 1, 0, 4, "Baos", "Calibri 35 bold")
                oListBao = []
                for b in baos.values():
                    oListBao.append(b['name'])
                baoSelected = StringVar()
                sauceSelected = StringVar()
                baoSelected.trace_add("write", baoChanged)
                sauceSelected.trace_add("write", sauceChanged)
                createText([OwnerAddItem], 2, 0, 1, "Bao: ")
                createDropdown([OwnerAddItem], 2, 1, 1, oListBao, baoSelected, "Calibri 15 bold", "No bao selected", 15)

                count = IntVar()
                createDropdown([OwnerAddItem], 2, 2, 1, [1, 2, 3, 4, 5], count, "Calibri 15 bold", 1, 5)

                createText([OwnerAddItem], 4, 0, 3, "Pickles: ", sticky='s', pady=10)
                createText([OwnerAddItem], 4, 2, 2, "Sauce: ", sticky='s', pady=10)

                createButton([OwnerAddItem], 8, 2, 1, "No Pickles", noPickles, "Calibri 15")

                note = StringVar()
                createText([OwnerAddItem], 11, 0, 1, "Notes: ", "Calibri 20", pady=10)
                createEntryBox([OwnerAddItem], 11, 1, 2, note, "Calibri 20", width=30, ipadx=30)


            case("bentos"):

                # Import here because idk honestly man I'm tired, probably circular import circumvention or something
                from orders import bentos, bentoSides, bentoPermittedSauces

                # Store the auto-generated UI elements and variables in a list so they can be easily accessed and cleared out
                UIElements = [[], [], []]
                vars = [[], [], []]
                
                # Sub-sub-function, runs when the bento in the dropdown is changed
                def bentoChanged(*args):
                    # Ignore if it is the init value
                    if('selected' in bentoSelected.get()): return
                    # Get the bento details in full
                    bentoStuff = bentos[keyFromVal(bentos, bentoSelected.get())]
                    bentoSide1 = bentoSides[keyFromVal(bentoSides, bentoSideSelected[0].get())]
                    bentoSide2 = bentoSides[keyFromVal(bentoSides, bentoSideSelected[1].get())]
                    # Re-used options list
                    oListMod = []
                    
                    # Destroy each auto-generated UI element
                    [e.destroy() for e in UIElements[0]]
                    
                    # Clear the lists
                    UIElements[0].clear()
                    vars[0].clear()

                    for i in range(0, len(bentoStuff['mod'])):
                        m = bentoStuff['mod'][i]
                        oListMod.clear()
                        vars[0].append(StringVar())
                        for v in modTypes[m['modType']].values():
                            oListMod.append(v)
                        UIElements[0].append(createText([OwnerAddItem], 5+i, 0, 1, f"{m['name']}: "))
                        UIElements[0].append(createDropdown([OwnerAddItem], 5+i, 1, 1, oListMod, vars[0][i], "Calibri 15 bold", modTypes[m['modType']][m['default']], 24))

                    oListBento.clear()
                    for a in bentoSides.values():
                        oListBento.append(a['name'])
                    UIElements[0].append(createText([OwnerAddItem], 10, 0, 2, "Side 1: "))
                    UIElements[0].append(createDropdown([OwnerAddItem], 11, 1, 1, oListBento, bentoSideSelected[0], "Calibri 15 bold", bentoSides[bentoStuff['side1'][0]]['name'], 24))
                    
                    oListBento.clear()
                    for a in bentoSides.values():
                        oListBento.append(a['name'])
                    UIElements[0].append(createText([OwnerAddItem], 20, 0, 2, "Side 2: "))
                    UIElements[0].append(createDropdown([OwnerAddItem], 21, 1, 1, oListBento, bentoSideSelected[1], "Calibri 15 bold", bentoSides[bentoStuff['side2'][0]]['name'], 24))

                    if(bentoStuff['sauce'] != -1):
                        oListSauce = []
                        for i in bentoPermittedSauces:
                            oListSauce.append(sauceDict[i])
                        UIElements[0].append(createDropdown([OwnerAddItem], 5, 2, 1, oListSauce, sauceSelected, "Calibri 15 bold", sauceDict[bentoStuff['sauce']], 24, 5))
                    
                    UIElements[0].append(createButton([OwnerAddItem], 26, 3, 4, "Add Item", addBento, "Calibri 20", pady=5))


                def bentoSideChanged(side: int):
                    s = side - 1
                    # Ignore if it is the init value
                    if(not(bentoSideSelected[s].get())): return
                    # Get the side details in full
                    bentoSideStuff = bentoSides[keyFromVal(bentoSides, bentoSideSelected[s].get())]
                    # Re-used options list
                    oListMod = []
                    
                    # Destroy each auto-generated UI element
                    [e.destroy() for e in UIElements[side]]
                    
                    # Clear the lists
                    UIElements[side].clear()
                    vars[side].clear()

                    for i in range(0, len(bentoSideStuff['mod'])):
                        m = bentoSideStuff['mod'][i]
                        oListMod.clear()
                        vars[side].append(StringVar())
                        for v in modTypes[m['modType']].values():
                            oListMod.append(v)
                        UIElements[side].append(createText([OwnerAddItem], (side*10)+1+i, 2, 1, f"{m['name']}: ", "Calibri 15"))
                        UIElements[side].append(createDropdown([OwnerAddItem], (side*10)+1+i, 3, 1, oListMod, vars[side][i], "Calibri 13 bold", modTypes[m['modType']][m['default']], 18))

                def addBento():
                    o.addItem(Bento(
                        count.get(),
                        keyFromVal(bentos, bentoSelected.get()),
                        [keyFromVal(modTypes[keyFromVal(modTypes, m.get())], m.get()) for m in vars[0]],

                        keyFromVal(bentoSides, bentoSideSelected[0].get()),
                        [keyFromVal(modTypes[keyFromVal(modTypes, m.get())], m.get()) for m in vars[1]],

                        keyFromVal(bentoSides, bentoSideSelected[1].get()),
                        [keyFromVal(modTypes[keyFromVal(modTypes, m.get())], m.get()) for m in vars[2]],
                        keyFromValPrecise(sauceDict, sauceSelected.get()),
                        1,
                        note.get()))
                    messagebox.showinfo("Success", "Successfully Added bento to order", parent=OwnerAddItem)
                    OwnerAddItem.destroy()

                createText([OwnerAddItem], 1, 0, 4, "Bentos", "Calibri 35 bold")
                oListBento = []
                for a in bentos.values():
                    oListBento.append(a['name'])
                bentoSelected = StringVar()
                bentoSideSelected = [StringVar(), StringVar()]
                bentoSelected.trace_add("write", bentoChanged)
                bentoSideSelected[0].trace_add("write", lambda x,y,z:bentoSideChanged(1))
                bentoSideSelected[1].trace_add("write", lambda x,y,z:bentoSideChanged(2))
                sauceSelected = StringVar()
                createText([OwnerAddItem], 2, 0, 1, "Bento: ")
                createDropdown([OwnerAddItem], 2, 1, 1, oListBento, bentoSelected, "Calibri 15 bold", "No bento selected", 24)

                count = IntVar()
                createDropdown([OwnerAddItem], 2, 2, 1, [1, 2, 3, 4, 5], count, "Calibri 15 bold", 1, 5)

                createText([OwnerAddItem], 4, 0, 2, "Modifiers: ", sticky='s', pady=10)
                createText([OwnerAddItem], 4, 2, 2, "Sauce: ", sticky='s', pady=10)

                note = StringVar()
                createText([OwnerAddItem], 26, 0, 1, "Notes: ", "Calibri 20", pady=10)
                createEntryBox([OwnerAddItem], 26, 1, 2, note, "Calibri 20", width=30, ipadx=30)

            case("classics"):
                # Import here because idk honestly man I'm tired, probably circular import circumvention or something
                from orders import classics, classicSides

                # Store the auto-generated UI elements and variables in a list so they can be easily accessed
                UIElements = [[], []]
                vars = [[], []]
                
                # Sub-sub-function, runs when the classic in the dropdown is changed
                def classicChanged(*args):
                    # Ignore if it is the init value
                    if('selected' in classicSelected.get()): return
                    # Get the classic details in full
                    classicStuff = classics[keyFromVal(classics, classicSelected.get())]
                    # Re-used options list
                    oListMod = []
                    
                    # Destroy each auto-generated UI element
                    [e.destroy() for e in UIElements[0]]
                    
                    # Clear the lists
                    UIElements[0].clear()
                    vars[0].clear()

                    for i in range(0, len(classicStuff['mod'])):
                        m = classicStuff['mod'][i]
                        oListMod.clear()
                        vars[0].append(StringVar())
                        for v in modTypes[m['modType']].values():
                            oListMod.append(v)
                        UIElements[0].append(createText([OwnerAddItem], 5+i, 0, 1, f"{m['name']}: "))
                        UIElements[0].append(createDropdown([OwnerAddItem], 5+i, 1, 1, oListMod, vars[0][i], "Calibri 15 bold", modTypes[m['modType']][m['default']], 24))

                    if(classicStuff['side'] != -1):
                        oListSides = []
                        for i in list(classicSides.values()):
                            oListSides.append(i['name'])
                        UIElements[0].append(createDropdown([OwnerAddItem], 5, 3, 1, oListSides, sideSelected, "Calibri 15 bold", classicSides[classicStuff['side']]['name'], 24, 5))
                    else:
                        sideSelected.set("")
                        UIElements[0].append(createText([OwnerAddItem], 5, 2, 1, "No side", "Calibri 20"))
                    
                    UIElements[0].append(createButton([OwnerAddItem], 13, 1, 4, "Add Item", addClassic, "Calibri 20", pady=5))


                def sideChanged(*args):
                    # Ignore if it is the init value
                    if(not(sideSelected.get())): return
                    # Get the side details in full
                    sideStuff = classicSides[keyFromVal(classicSides, sideSelected.get())]
                    # Re-used options list
                    oListMod = []
                    
                    # Destroy each auto-generated UI element
                    [e.destroy() for e in UIElements[1]]
                    
                    # Clear the lists
                    UIElements[1].clear()
                    vars[1].clear()

                    for i in range(0, len(sideStuff['mod'])):
                        m = sideStuff['mod'][i]
                        oListMod.clear()
                        vars[1].append(StringVar())
                        for v in modTypes[m['modType']].values():
                            oListMod.append(v)
                        UIElements[1].append(createText([OwnerAddItem], 6+i, 2, 1, f"{m['name']}: ", "Calibri 15"))
                        UIElements[1].append(createDropdown([OwnerAddItem], 6+i, 3, 1, oListMod, vars[1][i], "Calibri 13 bold", modTypes[m['modType']][m['default']], 18))

                def addClassic():
                    o.addItem(Classic(
                        count.get(),
                        keyFromVal(classics, classicSelected.get()),
                        [keyFromVal(modTypes[keyFromVal(modTypes, m.get())], m.get()) for m in vars[0]],
                        keyFromVal(classicSides, sideSelected.get()),
                        [keyFromVal(modTypes[keyFromVal(modTypes, m.get())], m.get()) for m in vars[1]],
                        note.get()))
                    messagebox.showinfo("Success", "Successfully added classic to order", parent=OwnerAddItem)
                    OwnerAddItem.destroy()

                createText([OwnerAddItem], 1, 0, 4, "Classics", "Calibri 35 bold")
                oListCla = []
                for a in classics.values():
                    oListCla.append(a['name'])
                classicSelected = StringVar()
                sideSelected = StringVar()
                classicSelected.trace_add("write", classicChanged)
                sideSelected.trace_add("write", sideChanged)
                createText([OwnerAddItem], 2, 0, 1, "Classic: ")
                createDropdown([OwnerAddItem], 2, 1, 1, oListCla, classicSelected, "Calibri 15 bold", "No classic selected", 24)

                count = IntVar()
                createDropdown([OwnerAddItem], 2, 2, 1, [1, 2, 3, 4, 5], count, "Calibri 15 bold", 1, 5)

                createText([OwnerAddItem], 4, 0, 2, "Modifiers: ", sticky='s', pady=10)
                createText([OwnerAddItem], 4, 2, 2, "Side: ", sticky='s', pady=10)

                note = StringVar()
                createText([OwnerAddItem], 12, 0, 1, "Notes: ", "Calibri 20", pady=10)
                createEntryBox([OwnerAddItem], 12, 1, 2, note, "Calibri 20", width=30, ipadx=30)
            case("sides"):

                # Import here because idk honestly man I'm tired, probably circular import circumvention or something
                from orders import sides

                # Store the auto-generated UI elements and variables in a list so they can be easily accessed
                UIElements = []
                vars = []

                def sideChanged(*args):
                    # Ignore if it is the init value
                    if('selected' in sideSelected.get()): return
                    # Get the side details in full
                    sideStuff = sides[keyFromVal(sides, sideSelected.get())]
                    # Re-used options list
                    oListMod = []
                    
                    # Destroy each auto-generated UI element
                    [e.destroy() for e in UIElements]
                    
                    # Clear the lists
                    UIElements.clear()
                    vars.clear()

                    for i in range(0, len(sideStuff['mod'])):
                        m = sideStuff['mod'][i]
                        oListMod.clear()
                        vars.append(StringVar())
                        for v in modTypes[m['modType']].values():
                            oListMod.append(v)
                        UIElements.append(createText([OwnerAddItem], 5+i, 0, 1, f"{m['name']}: ", "Calibri 15"))
                        UIElements.append(createDropdown([OwnerAddItem], 5+i, 1, 1, oListMod, vars[i], "Calibri 13 bold", modTypes[m['modType']][m['default']], 18))

                    UIElements.append(createButton([OwnerAddItem], 13, 1, 4, "Add Item", addSide, "Calibri 20", pady=5))

                def addSide():
                    o.addItem(Side(
                        count.get(),
                        keyFromVal(sides, sideSelected.get()),
                        [keyFromVal(modTypes[keyFromVal(modTypes, m.get())], m.get()) for m in vars],
                        note.get()))
                    messagebox.showinfo("Success", "Successfully added side to order", parent=OwnerAddItem)
                    OwnerAddItem.destroy()

                createText([OwnerAddItem], 1, 0, 4, "Sides", "Calibri 35 bold")
                oListSid = []
                for a in sides.values():
                    oListSid.append(a['name'])
                sideSelected = StringVar()
                sideSelected.trace_add("write", sideChanged)
                createText([OwnerAddItem], 2, 0, 1, "Side: ")
                createDropdown([OwnerAddItem], 2, 1, 1, oListSid, sideSelected, "Calibri 15 bold", "No side selected", 24)

                count = IntVar()
                createDropdown([OwnerAddItem], 2, 2, 1, [1, 2, 3, 4, 5], count, "Calibri 15 bold", 1, 5)

                createText([OwnerAddItem], 4, 0, 2, "Modifiers: ", sticky='s', pady=10)

                note = StringVar()
                createText([OwnerAddItem], 12, 0, 1, "Notes: ", "Calibri 20", pady=10)
                createEntryBox([OwnerAddItem], 12, 1, 2, note, "Calibri 20", width=30, ipadx=30)



#############################################################################################################################
#############################################################################################################################
#############################################################################################################################


#


#############################################################################################################################
#############################################################################################################################
#############################################################################################################################

def createOwnerViewOrdersToplevel():
    OwnerViewOrders = Toplevel(root, bg='#000000')
    OwnerViewOrders.protocol("WM_DELETE_WINDOW", lambda: closedWindow(OwnerViewOrders))
    OwnerViewOrders.geometry('1500x500')
    createText([OwnerViewOrders], 1, 0, 4, "View Orders", "Calibri 35 bold")

    columns = ("oid", "cid", "data", "com", "paid", "ptime", "pktime")
    headings = ("Order ID", "Customer ID", "Order Data", "Completed?", "Paid?", "Placement Time", "Pickup Time")
    widths = (100, 100, 100, 80, 80, 180, 180)

    treeview = ttk.Treeview(OwnerViewOrders, columns=columns, show='headings', padding=1, selectmode='browse')
    treeview.grid(row = 3, column = 0, columnspan=5, padx=10)

    orderTreeview = ttk.Treeview(OwnerViewOrders, padding=3, columns=["title", "price"], show='headings')
    orderTreeview.grid(row=3, column=5, padx=10, columnspan=2)
    orderTreeview.column("title", anchor='w', width=400)
    orderTreeview.column("price", anchor='w', width=100)
    orderTreeview.heading("title", text='No order selected', anchor='center')

    def populate(filter=""):
        treeview.delete(*treeview.get_children())
        
        cursor.execute(f"""SELECT * FROM orders {filter}""")
        data = cursor.fetchall()

        for d in data:
            
            itemCount = sum([sum([i['count'] for i in v]) for v in json.loads(d[2]).values()])
            treeview.insert("", END, values=[
                d[0], 
                d[1], 
                f"{itemCount} item{'s' if itemCount > 1 else ''}", 
                "Yes" if d[3] else "No", "Yes" if d[4] else "No", 
                strftime("%a, %d %b %Y %H:%M:%S", time.localtime(d[5] + 1704067200)), 
                strftime("%a, %d %b %Y %H:%M:%S", time.localtime(d[6] + 1704067200))
            ])
            treeview.bind('<<TreeviewSelect>>', lambda event: orderSelected(event, orderTreeview))
        for c, h, w in zip(columns, headings, widths):
            treeview.column(c, anchor='center', width=w)
            treeview.heading(c, text=h, anchor='center')

    def markComplete():
        global selectedOrder
        if(not(selectedOrder)):
            messagebox.showerror("Error", "No order selected", parent=treeview)
            return
        cursor.execute(f"UPDATE orders SET complete = {0 if selectedOrder[3] == 'Yes' else 1} WHERE orderID = {selectedOrder[0]};")
        db.commit()
        
    def markPaid():
        global selectedOrder
        if(not(selectedOrder)):
            messagebox.showerror("Error", "No order selected", parent=treeview)
            return
        cursor.execute(f"UPDATE orders SET paid = {0 if selectedOrder[4] == 'Yes' else 1} WHERE orderID = {selectedOrder[0]};")
        db.commit()

    buttonFont = 'Calibri 18'
    createText([OwnerViewOrders], 4, 0, 1, "Filter:", pady=8)
    createButton([OwnerViewOrders], 4, 1, 1, "All", command=lambda: populate(""), font=buttonFont)
    createButton([OwnerViewOrders], 4, 2, 1, "Active", command=lambda: populate("WHERE complete = 0"), font=buttonFont)
    createButton([OwnerViewOrders], 4, 3, 1, "Complete", command=lambda: populate("WHERE complete = 1"), font=buttonFont)
    createButton([OwnerViewOrders], 4, 4, 1, "Outstanding", command=lambda: populate("WHERE complete = 1 AND paid = 0"), font=buttonFont)

    createButton([OwnerViewOrders], 4, 5, 1, "Mark Complete", command=markComplete, font=buttonFont)
    createButton([OwnerViewOrders], 4, 6, 1, "Mark Paid", command=markPaid, font=buttonFont)
    createText([OwnerViewOrders], 5, 5, 2, "Make sure to check that you have the right\norder and to refresh after marking orders!", font="Calibri 10")
    
    
    populate()

    createButton([OwnerViewOrders], 5, 2, 2, "Generate Order Receipt", command=lambda:createOrderReceipt(OwnerViewOrders), font=buttonFont, width=20)


def createOrderReceipt(parentTV, orderData = {}):
    if(orderData == {}):
        global selectedOrder
        if(not(selectedOrder)):
            messagebox.showerror("Error", "No order selected", parent=parentTV)
            return
    
    from orders import modTypes, sauceDict, picklesDict, appetisers, baos, bentos, bentoSides, classics, classicSides, sides
    
    try:
        if(orderData == {}):
            orderNum = selectedOrder[0]
            open(f'documents/orderReceipts/{orderNum}.docx')
        else:
            orderNum = 'NA - Temporary receipt'
            raise FileNotFoundError
    except FileNotFoundError:
        doc = Document()
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.1)
            section.bottom_margin = Inches(0.1)
            section.left_margin = Inches(0.1)
            section.right_margin = Inches(0.1)

        doc.add_heading(f'Order {orderNum}', 0)

        doc.add_picture('images/bb_logo_b.png', width=Inches(2))

        if(orderData == {}):
            cursor.execute(f"""SELECT *  FROM orders WHERE orderID = {orderNum}""")
            orderData = json.loads(str(cursor.fetchone()[2]).removeprefix('b\'').removesuffix('\''))

        doc.add_heading('Appetisers', level=1)

        items = []
        for d in orderData['appetisers']:
            mods = []

            for i in range(0, len(d['details'][1])):
                modMod = modTypes[appetisers[d['details'][0]]['mod'][i]['modType']][d['details'][1][i]]
                modName = appetisers[d['details'][0]]['mod'][i]['name']
                if(modName == "Well done" and modMod == "Not"): continue
                if(modName == "Hot" and modMod == "Normal"): continue
                if(modName == "Well done"): mods.append(f'{modName}'); continue
                if(modName == "Hot"): mods.append(f'{modMod}'); continue
                if(modName == "Lemon" and modMod == "Yes"): continue
                mods.append(f'{modMod} {modName}')
            items.append((
                f'{d['count']} {appetisers[d['details'][0]]['name']}', 
                f'{sauceDict[d['details'][2]]}', 
                mods
            ))

        table = doc.add_table(rows=1, cols=3)
        header = table.rows[0].cells
        header[0].text = 'Qty&Name'
        header[1].text = 'Sauce'
        header[2].text = 'Mod.'
        for qname, sauce, mod in items:
            row = table.add_row().cells
            row[0].text = qname
            row[1].text = sauce
            for m in mod:
                row[2].text += f'{m}\n'
            row[2].text = row[2].text.removesuffix('\n')


        doc.add_heading('Baos', level=1)
        
        items = []
        for d in orderData['baos']:
            items.append((
                f'{d['count']} {baos[d['details'][0]]['name']}', 
                f'{f"{modTypes[5][d['details'][2]]} " if d['details'][1] != 0 else ""}{sauceDict[d['details'][1]]}', 
                [f'{modTypes[1][d['details'][3][i]]} {picklesDict[i + 1]}' for i in range(0, 5)]
            ))

        table = doc.add_table(rows=1, cols=3)
        header = table.rows[0].cells
        header[0].text = 'Qty&Name'
        header[1].text = 'Sauce'
        header[2].text = 'Pickles'
        for qname, sauce, mod in items:
            row = table.add_row().cells
            row[0].text = qname
            row[1].text = sauce
            for m in mod:
                row[2].text += f'{m}\n'
            row[2].text = row[2].text.removesuffix('\n')
        
        
        doc.add_heading('Bentos', level=1)

        items = []
        for d in orderData['bentos']:
            mods = []
            s1mods = []
            s2mods = []

            for i in range(0, len(d['details'][1])):
                modMod = modTypes[bentos[d['details'][0]]['mod'][i]['modType']][d['details'][1][i]]
                modName = bentos[d['details'][0]]['mod'][i]['name']
                if(modName == "Hot" and modMod == "Normal"): continue
                if(modName == "Hot"): mods.append(f'{modMod}'); continue
                mods.append(f'{modMod} {modName}')

            for i in range(0, len(d['details'][3])):
                s1modMod = modTypes[bentoSides[d['details'][2]]['mod'][i]['modType']][d['details'][3][i]]
                s1modName = bentoSides[d['details'][2]]['mod'][i]['name']
                s1mods.append(f'{s1modMod} {s1modName}')

            for i in range(0, len(d['details'][5])):
                s2modMod = modTypes[bentoSides[d['details'][4]]['mod'][i]['modType']][d['details'][5][i]]
                s2modName = bentoSides[d['details'][4]]['mod'][i]['name']
                s2mods.append(f'{s2modMod} {s2modName}')

            items.append((
                f'{d['count']} {bentos[d['details'][0]]['name']}', 
                f'{sauceDict[d['details'][6]]}',
                '',
                mods,
                f'{bentoSides[d['details'][2]]['name']}',
                s1mods,
                f'{bentoSides[d['details'][4]]['name']}',
                s2mods,
            ))

        table = doc.add_table(rows=1, cols=7)
        header = table.rows[0].cells
        header[0].text = 'Qty&Name'
        header[1].text = 'Sauce'
        header[2].text = 'Mod.'
        header[3].text = 'Side 1'
        header[4].text = 'S1 Mod.'
        header[5].text = 'Side 2'
        header[6].text = 'S2 Mod.'
        for qname, sauce, sauceMod, mod, s1, s1m, s2, s2m in items:
            row = table.add_row().cells
            row[0].text = qname
            row[1].text = f'{sauce}{f" ({sauceMod})" if sauceMod else ""}'
            for m in mod:
                row[2].text += f'{m}\n'
            row[2].text = row[2].text.removesuffix('\n')
            row[3].text = s1
            for m1 in s1m:
                row[4].text += f'{m1}\n'
            row[4].text = row[4].text.removesuffix('\n')
            row[5].text = s2
            for m2 in s2m:
                row[6].text += f'{m2}\n'
            row[6].text = row[6].text.removesuffix('\n')

        
        doc.add_heading('Classics', level=1)
        
        items = []
        for d in orderData['classics']:
            mods = []
            sMods = []

            for i in range(0, len(d['details'][1])):
                modMod = modTypes[classics[d['details'][0]]['mod'][i]['modType']][d['details'][1][i]]
                modName = classics[d['details'][0]]['mod'][i]['name']
                if(modName == "Hot" and modMod == "Normal"): continue
                if(modName == "Hot"): mods.append(f'{modMod}'); continue
                mods.append(f'{modMod} {modName}')

            for i in range(0, len(d['details'][3])):
                sModMod = modTypes[classicSides[d['details'][2]]['mod'][i]['modType']][d['details'][3][i]]
                sModName = classicSides[d['details'][2]]['mod'][i]['name']
                sMods.append(f'{sModMod} {sModName}')

            items.append((
                f'{d['count']} {classics[d['details'][0]]['name']}', 
                mods,
                f'{classicSides[d['details'][2]]['name']}',
                sMods,
            ))

        table = doc.add_table(rows=1, cols=4)
        header = table.rows[0].cells
        header[0].text = 'Qty&Name'
        header[1].text = 'Mod.'
        header[2].text = 'Side'
        header[3].text = 'S Mod.'
        for qname, mod, s, sm in items:
            row = table.add_row().cells
            row[0].text = qname
            for m in mod:
                row[1].text += f'{m}\n'
            row[1].text = row[1].text.removesuffix('\n')
            row[2].text = s
            for m in sm:
                row[3].text += f'{m}\n'
            row[3].text = row[3].text.removesuffix('\n')
        


        doc.add_heading('Sides', level=1)

        items = []
        for d in orderData['sides']:
            mods = []
            sMods = []

            for i in range(0, len(d['details'][1])):
                modMod = modTypes[sides[d['details'][0]]['mod'][i]['modType']][d['details'][1][i]]
                modName = sides[d['details'][0]]['mod'][i]['name']
                if(modName == "Hot" and modMod == "Normal"): continue
                if(modName == "Onion" and modMod == "Yes"): mods.append(f'Plus {modName}'); continue
                if(modName == "Onion" and modMod == "No"): continue
                mods.append(f'{modMod} {modName}')

            items.append((
                f'{d['count']} {sides[d['details'][0]]['name']}', 
                mods
            ))

        table = doc.add_table(rows=1, cols=2)
        header = table.rows[0].cells
        header[0].text = 'Qty&Name'
        header[1].text = 'Mod.'
        for qname, mod in items:
            row = table.add_row().cells
            row[0].text = qname
            for m in mod:
                row[1].text += f'{m}\n'
            row[1].text = row[1].text.removesuffix('\n')
        
        doc.add_page_break()

        doc.save(f'documents/orderReceipts/{orderNum}.docx')

    finally:
        if os.system(f'start documents/orderReceipts/"{orderNum}".docx') != 0:
            messagebox.showerror("Error", "Cannot open receipt, file already open!", parent=parentTV)

#############################################################################################################################
#############################################################################################################################
#############################################################################################################################

def createOwnerManageCustomersToplevel():
    OwnerManageCustomers = Toplevel(root, bg='#000000')
    OwnerManageCustomers.geometry('1230x500')
    createText([OwnerManageCustomers], 1, 0, 4, "Manage Customers", "Calibri 35 bold")

    columns = ("cid", "fn", "ln", "pn", "em", "pw")
    headings = ("Customer ID", "First Name", "Last Name", "Phone Number", "Email", "Password")

    treeview = ttk.Treeview(OwnerManageCustomers, columns=columns, show='headings', padding=1, selectmode='browse')
    cursor.execute(f"SELECT * FROM customers")
    data = cursor.fetchall()
    for d in data:
        treeview.insert("", END, values=d)
    for c, h in zip(columns, headings):
        treeview.heading(c, text=h)
    treeview.grid(row = 3, column = 0, padx=10)
    treeview.bind('<<TreeviewSelect>>', lambda event: customerSelected(event, treeview))

def customerSelected(event, tree):
    selected_items = event.widget.selection()
    if selected_items:
        item = selected_items[0]
        record = event.widget.item(item)['values']
        createOwnerModerationToplevel(record)

#############################################################################################################################

def createOwnerModerationToplevel(customer: list):
    OwnerManageIndividual = Toplevel(root, bg='#000000')
    createText([OwnerManageIndividual], 1, 0, 4, "Manage Customer", "Calibri 35 bold")
    createText([OwnerManageIndividual], 2, 0, 4, f"{customer[0]}, {customer[1]} {customer[2]}", "Calibri 20 bold")
    createButton([OwnerManageIndividual], 3, 0, 4, "View Order History", lambda:createOwnerViewCustomerOrdersToplevel(customer), "Calibri 18", ipadx=15)

#############################################################################################################################

def createOwnerViewCustomerOrdersToplevel(customer):
    OwnerViewCustomerOrders= Toplevel(root, bg='#000000')
    OwnerViewCustomerOrders.protocol("WM_DELETE_WINDOW", lambda: closedWindow(OwnerViewCustomerOrders))
    OwnerViewCustomerOrders.geometry("1500x500")
    createText([OwnerViewCustomerOrders], 1, 0, 4, "Order History", "Calibri 35 bold")
    createText([OwnerViewCustomerOrders], 2, 0, 4, f"{customer[0]}, {customer[1]} {customer[2]}", "Calibri 20 bold")

    data = retrieveCustomerOrders(customer[0])

    columns = ("oid", "data", "com", "paid", "ptime", "pktime")
    headings = ("Order ID", "Order Data", "Completed?", "Paid?", "Placement Time", "Pickup Time")
    widths = (100, 100, 80, 80, 180, 180)
    treeview = ttk.Treeview(OwnerViewCustomerOrders, columns=columns, show='headings', padding=1, selectmode='browse')

    orderTreeview = ttk.Treeview(OwnerViewCustomerOrders, padding=3, columns=["title", "price"], show='headings')
    orderTreeview.grid(row=3, column=1, padx=10, columnspan=2)
    orderTreeview.column("title", anchor='w', width=400)
    orderTreeview.column("price", anchor='w', width=100)
    orderTreeview.heading("title", text='No order selected', anchor='center')

    for d in data:
        itemCount = sum([sum([i['count'] for i in v]) for v in json.loads(d[2]).values()])
        treeview.insert("", END, values=[d[0], f"{itemCount} item{'s' if itemCount > 1 else ''}", "Yes" if d[3] else "No", "Yes" if d[4] else "No",  strftime("%a, %d %b %Y %H:%M:%S", time.localtime(d[5] + 1704067200)), strftime("%a, %d %b %Y %H:%M:%S", time.localtime(d[6] + 1704067200))])
        treeview.bind('<<TreeviewSelect>>', lambda event: orderSelected(event, orderTreeview))
    for c, h, w in zip(columns, headings, widths):
        treeview.column(c, anchor='center', width=w)
        treeview.heading(c, text=h, anchor='center')
    treeview.grid(row = 3, column = 0, padx=10)

#############################################################################################################################

def orderSelected(event, tree):
    selectedItems = event.widget.selection()
    if selectedItems:
        item = selectedItems[0]
        record = event.widget.item(item)['values']
        global selectedOrder
        selectedOrder = record
        cursor.execute(f"""SELECT *  FROM orders WHERE orderID = {record[0]}""")
        orderData = cursor.fetchone()[2]

        tree.heading("title", text=f'Order {record[0]}', anchor='center')
        tree.heading("price", text='Price:', anchor='center')

        tree.delete(*tree.get_children())

        totalPrice = 0.00

        for k, v in json.loads(orderData).items():
            order = tree.insert("", END, values=[f"[{sum([i['count'] for i in v])}] {k.title()}"])
            for jtem in v:
                match(k):
                    case("appetisers"):
                        appItem = Appetiser(jtem['count'], *jtem['details'], jtem['note'])
                        text = appItem.outputPretty()
                        appetisers = tree.insert(order, END, values=[f"    {text['title']}", floatToPrice(text['price'])])
                        if(appItem.note):
                            tree.insert(appetisers, END, values=[f"            Note: {appItem.note}"])
                        if(text['sauce']):
                            tree.insert(appetisers, END, values=[f"            {text['sauce']}"])
                        totalPrice += float(text['price'])
                    case("baos"):
                        baoItem = Bao(jtem['count'], *jtem['details'], jtem['note'])
                        text = baoItem.outputPretty()
                        baos = tree.insert(order, END, values=[f"    {text['title']}", floatToPrice(text['price'])])
                        if(baoItem.note):
                            tree.insert(baos, END, values=[f"            Note: {baoItem.note}"])
                        tree.insert(baos, END, values=[f"            {text['sauce']}"])
                        for p in text['pickles']:
                            tree.insert(baos, END, values=[f"            {p}"])
                        totalPrice += float(text['price'])
                    case("bentos"):
                        bentoItem = Bento(jtem['count'], *jtem['details'], jtem['note'])
                        text = bentoItem.outputPretty()
                        bentos = tree.insert(order, END, values=[f"    {text['title']}", floatToPrice(text['price'])])
                        if(bentoItem.note):
                            tree.insert(bentos, END, values=[f"            Note: {bentoItem.note}"])
                        side1 = tree.insert(bentos, END, values=[f"            {text['side1']}"])
                        for m in text['side1Mod']:
                            tree.insert(side1, END, values=[f"                  {m}"])
                        side2 = tree.insert(bentos, END, values=[f"            {text['side2']}"])
                        for m in text['side2Mod']:
                            tree.insert(side2, END, values=[f"                  {m}"])
                        totalPrice += float(text['price'])
                    case("classics"):
                        classicItem = Classic(jtem['count'], *jtem['details'], jtem['note'])
                        text = classicItem.outputPretty()
                        classics = tree.insert(order, END, values=[f"    {text['title']}", floatToPrice(text['price'])])
                        if(classicItem.note):
                            tree.insert(classics, END, values=[f"            Note: {classicItem.note}"])
                        side = tree.insert(classics, END, values=[f"            {text['side']}"])
                        for m in text['sideMod']:
                            tree.insert(side, END, values=[f"                   {m}"])
                        totalPrice += float(text['price'])
                    case("sides"):
                        sideItem = Side(jtem['count'], *jtem['details'], jtem['note'])
                        text = sideItem.outputPretty()
                        sides = tree.insert(order, END, values=[f"    {text['title']}", floatToPrice(text['price'])])
                        if(sideItem.note):
                            tree.insert(sides, END, values=[f"            Note: {sideItem.note}"])
                        for m in text['sideMod']:
                            tree.insert(sides, END, values=[f"            {m}"])
                        totalPrice += float(text['price'])

        tree.insert("", END, values=[""])
        tree.insert("", END, values=[f"Total: ", floatToPrice(totalPrice)])

#############################################################################################################################
#############################################################################################################################
#############################################################################################################################

def createOwnerManageEmployeesToplevel():
    OwnerManageEmployees = Toplevel(root, bg='#000000')
    OwnerManageEmployees.geometry('1230x500')
    createText([OwnerManageEmployees], 1, 0, 4, "Manage Employees", "Calibri 35 bold")

    columns = ("cid", "fn", "ln", "ak")
    headings = ("Employee ID", "First Name", "Last Name", "Access Key")

    treeview = ttk.Treeview(OwnerManageEmployees, columns=columns, show='headings', padding=1, selectmode='browse')
    cursor.execute(f"SELECT * FROM employees")
    data = cursor.fetchall()
    for d in data:
        treeview.insert("", END, values=d)
    for c, h in zip(columns, headings):
        treeview.heading(c, text=h)
    treeview.grid(row = 3, column = 0, padx=10)
    treeview.bind('<<TreeviewSelect>>', lambda event: employeeSelected(event, treeview))

def employeeSelected(event, tree):
    selected_items = event.widget.selection()
    if selected_items:
        item = selected_items[0]
        record = event.widget.item(item)['values']
        # do some stuff

#############################################################################################################################
#############################################################################################################################
#############################################################################################################################
        
def createOwnerReportsToplevel():

    # Import item details
    from orders import appetisers, baos, bentos, classics, sides

    # Create toplevel
    OwnerReports = Toplevel(root, bg='#000000')
    OwnerReports.protocol("WM_DELETE_WINDOW", lambda: closedWindow(OwnerReports))
    OwnerReports.geometry("1500x500")
    createText([OwnerReports], 1, 0, 4, "Reports", "Calibri 35 bold")

    oList = ["All Time", "Current Month", "Current Year"]
    oListMonth = ["Whole Year"]
    UIElements = [[], [], []]

    for y in range(2024, int(strftime("%Y")) + 1):
        oList.append(strftime("%Y", time.struct_time([y, 1, 1, 0, 0, 0, 0, 1, 0])))

    selectedPeriod = StringVar()
    selectedMonth = StringVar()
    selectedItemCat = StringVar()

    createDropdown([OwnerReports], 2, 0, 2, oList, selectedPeriod, "Calibri 20", "All Time")

    [e.destroy() for e in UIElements[1]]
    UIElements[1].clear()
    oListItemCats = ["Appetisers", "Baos", "Bentos", "Classics", "Sides"]
    UIElements[1].append(createDropdown([OwnerReports], 5, 2, 1, oListItemCats, selectedItemCat, "Calibri 20", "Appetisers"))
    UIElements[1].append(createText([OwnerReports], 6, 2, 1, "Note: Generating an item trend graph for a specific\nmonth won't yield particularly useful results", "Calibri 15"))

    def periodChanged(*args):
        [e.destroy() for e in UIElements[0]]
        UIElements[0].clear()
        if(selectedPeriod.get() in ["All Time", "Current Month", "Current Year"]):
            return
        oListMonth = ["Whole Year"]
        if(int(strftime("%Y")) > int(selectedPeriod.get())): r = 13
        else: r = int(strftime("%m"))
        for m in range(1, r):
            oListMonth.append(strftime("%B", time.struct_time([y, m, 1, 0, 0, 0, 0, 1, 0])))
        UIElements[0].append(createDropdown([OwnerReports], 2, 2, 2, oListMonth, selectedMonth, "Calibri 20", "Whole Year"))


    cat = [{}]
    def itemCatChanged(*args):
        [e.destroy() for e in UIElements[2]]
        UIElements[2].clear()
        match(selectedItemCat.get()):
            case("Appetisers"):
                cat[0] = appetisers
            case("Baos"):
                cat[0] = baos
            case("Bentos"):
                cat[0] = bentos
            case("Classics"):
                cat[0] = classics
            case("Sides"):
                cat[0] = sides
    
    selectedPeriod.trace_add("write", periodChanged)
    selectedItemCat.trace_add("write", itemCatChanged)

    periodChanged()
    itemCatChanged()
    
    # # Document generation for sold food items
    # def generateSoldFoodItemsReport(timeMin = 0, timeMax = 1704067200):
    #     doc = Document()
    #     doc.add_heading('Sold Food Items', 0)

    #     doc.add_picture('images/bb_logo_b.png', width=Inches(2))

    #     match(selectedPeriod.get()):
    #         case("Current Month"):
    #             doc.add_paragraph(f"For the month of {strftime("%B")}")
    #             doc.add_paragraph("As of the " + ordinal(int(strftime("%d"))))
    #         case("Current Year"):
    #             doc.add_paragraph(f"For the year of {strftime("%Y")}")
    #             doc.add_paragraph(f"As of the {ordinal(int(strftime("%d")))} of {strftime("%B")}")
    #         case("All Time"):
    #             doc.add_paragraph(f"All current data")
    #         case(_):
    #             doc.add_paragraph(f"For the month of {strftime("%B, %Y", time.localtime(timeMin))}")
            
    #     cursor.execute(f"""SELECT orderData FROM orders WHERE placementTime > {timeMin} AND placementTime < {timeMax}""")
    #     data = cursor.fetchall()
    #     app = [[o['name'], 0] for o in appetisers.values()]
    #     bao = [[o['name'], 0] for o in baos.values()]
    #     ben = [[o['name'], 0] for o in bentos.values()]
    #     #cla = [[o['name'], 0] for o in classics.values()]
    #     #sid = [[o['name'], 0] for o in sides.values()]
        
    #     for dat in data:
    #         d = json.loads(str(dat[0]).removeprefix('b\'').removesuffix('\''))
    #         for a in d['appetisers']:
    #             app[a['details'][0] - 1][1] += a['count']
    #         for b in d['baos']:
    #             bao[b['details'][0] - 1][1] += b['count']
    #         for b in d['bentos']:
    #             ben[b['details'][0] - 1][1] += b['count']
    #         for c in d['classics']:
    #             pass
    #         for s in d['sides']:
    #             pass

    #     records = (["Appetisers", app], ["Baos", bao], ["Bentos", ben])#, ["Classics", cla], ["Sides", sid])

    #     for r in records:
    #         doc.add_heading(r[0], level=1)
    #         table = doc.add_table(rows=1, cols=2)
    #         header = table.rows[0].cells
    #         header[0].text = 'Name'
    #         header[1].text = 'Qty Sold'
    #         total = 0
    #         for name, qty in r[1]:
    #             row = table.add_row().cells
    #             row[0].text = name
    #             row[1].text = str(qty)
    #             total += qty
    #         doc.add_page_break()
    #         row = table.add_row().cells
    #         row[0].text = "Total sold"
    #         row[1].text = str(total)

    #     doc.add_page_break()

    #     try:
    #         doc.save('documents/report.docx')
    #         os.system('start documents/report.docx')
    #     except PermissionError:
    #         messagebox.showerror("Error", "File already open!\nPlease close the active report file to update the details", parent=OwnerReports)
    
    def generateSoldFoodItemsGraph():
        title, timeRange = graphTitle()

        title = f"{selectedItemCat.get()}\n{title}"
                   
        graphData = {}

        for i in list(cat[0].values()):
            graphData.update({i['name']: 0})

        curTimeCheck = timeRange[0]
        while curTimeCheck < timeRange[1]:
            curTime = time.localtime(curTimeCheck)
            timeNext = curTime
            timeNext = time.struct_time([timeNext.tm_year, timeNext.tm_mon + 1, *timeNext[2:9]])
            curTimeCheck = time.mktime(timeNext)

            cursor.execute(f"""SELECT orderData FROM orders WHERE placementTime > {time.mktime(curTime)} AND placementTime < {time.mktime(timeNext)}""")
            data = cursor.fetchall()
        
            for dat in data:
                d = json.loads(str(dat[0]).removeprefix('b\'').removesuffix('\''))[selectedItemCat.get().lower()]
                for i in d:
                    try:
                        graphData[cat[0][i['details'][0]]['name']] += i['count']
                    except IndexError:
                        graphData[cat[0][i['details'][0]]['name']].append(i['count'])
                    
        plotBarChart(title, "Item", "Qty Sold", graphData)

    def generateItemTrendGraph():
        title, timeRange = graphTitle()

        title = f"{selectedItemCat.get()}\n{title}"
                   
        graphData = {}
        xLabels = []

        for i in list(cat[0].values()):
            graphData.update({i['name']: [0 for i in range(0, 12)]})

        curTimeCheck = timeRange[0]
        m = 0
        while curTimeCheck < timeRange[1]:
            curTime = time.localtime(curTimeCheck)
            xLabels.append(f'{strftime("%b, %Y", time.localtime(curTimeCheck + 1704067200))}')
            timeNext = curTime
            temp = [timeNext.tm_year, timeNext.tm_mon + 1, 2, 0, 0, 0, *timeNext[6:9]]
            timeNext = time.struct_time(temp)
            curTimeCheck = time.mktime(timeNext)

            cursor.execute(f"""SELECT orderData FROM orders WHERE placementTime > {time.mktime(curTime)} AND placementTime < {time.mktime(timeNext)}""")
            data = cursor.fetchall()

            for dat in data:
                d = json.loads(str(dat[0]).removeprefix('b\'').removesuffix('\''))[selectedItemCat.get().lower()]
                for i in d:
                    graphData[cat[0][i['details'][0]]['name']][m] += i['count']
            m = m + 1
        plotLineChart(title, "Month, Year", "Qty Sold", graphData, xLabels)

    def plotBarChart(graphName: str, xName: str, yName: str, data: dict):
        import matplotlib.pyplot as mpl
        
        fig = mpl.figure(figsize=(14.0, 8.0))
        graph = fig.subplots()
        #graph = mpl.figure(figsize=(8.0, 4.5))

        count = 0
        for k, v in data.items():
            if(v) != 0:
                graph.bar(truncateText(k, math.floor(25-(1.3 * len(data.items())))), v, 0.9).set_label(k)
                count += 1
        
        graph.set_xlabel(xName)
        graph.set_ylabel(yName)
        graph.set_title(graphName)
        graph.set_xlim(0.6)
        graph.axis([-0.6, count, 0, max([i for i in data.values()]) + 1])

        try:
            graph.set_yticks([y for y in range(0, max([i for i in data.values()]) + 1)])
            # Show the legend
            graph.legend()
            
            # Show the plot in a separate window
            mpl.show()
        except ValueError:
            messagebox.showerror("Error", "No items in category for given time range", parent=OwnerReports)

    def plotLineChart(graphName: str, xName: str, yName: str, data: dict, xLabels: list[str]):
        import matplotlib.pyplot as mpl
        
        fig = mpl.figure(figsize=(14.0, 8.0))
        graph = fig.subplots()


        count = (len(list(data.values())))/-150
        for k, v in data.items():
            mpl.plot([i for i in range(0, len(v))], [adj + count for adj in v], label=k)
            count += 1/75
        
        graph.set_xlabel(xName)
        graph.set_ylabel(yName)
        graph.set_title(graphName)
        graph.set_xlim(0.6)
        graph.axis([0, count, 0, max([max([0, *i]) for i in data.values()]) + 1])

        try:
            graph.set_xticks([i for i in range(0, len(xLabels))], xLabels)
            graph.set_yticks([y for y in range(0, max([max([0, *i]) for i in data.values()]) + 1)])
            # Show the legend
            graph.legend()
            
            # Show the plot in a separate window
            mpl.show()
        except ValueError:
            messagebox.showerror("Error", "No items in category for given time range", parent=OwnerReports)
    
    def graphTitle():
        monYr = time.localtime()
        timeRange = [0, time.time() - 1704067200]
        match(selectedPeriod.get()):
            case("Current Month"):
                timeRange = [monthToTime(monYr.tm_mon, monYr.tm_year), monthToTime(monYr.tm_mon + 1, monYr.tm_year)]
                title = f"For the month of {strftime("%B")}\nAs of the {ordinal(int(strftime("%d")))}"
            case("Current Year"):
                timeRange = [monthToTime(1, monYr.tm_year), monthToTime(1, time.localtime().tm_year + 1)]
                title = f"For the year of {strftime("%Y")}\nAs of the {ordinal(int(strftime("%d")))} of {strftime("%B")}"
            case("All Time"):
                title = f"All current data"
            case(_):
                if(selectedMonth.get() == "Whole Year"):
                    timeRange = [monthToTime(1, selectedPeriod.get()), monthToTime(1, int(selectedPeriod.get()) + 1) - 86400]
                    if(int(selectedPeriod.get()) == int(strftime("%Y"))):
                        title = f"For the year of {strftime("%Y", time.localtime(timeRange[0] + 1704067200))}\nAs of the {ordinal(int(strftime("%d")))} of {strftime("%B")}"
                    else:
                        title = f"For the year of {strftime("%Y", time.localtime(timeRange[0] + 1704067200))}"
                else:
                    timeRange = [monthToTime(indexFromVal(months, selectedMonth.get()) + 1, selectedPeriod.get()), monthToTime(indexFromVal(months, selectedMonth.get()) + 2, selectedPeriod.get())]
                    title = f"For the month of {strftime("%B, %Y", time.localtime(timeRange[0] + 1704067200))}"
        return title, timeRange
        


    # createButton([OwnerReports], 4, 0, 2, "Sold Food Items - Quantities", lambda:generateSoldFoodItemsReport(), "Calibri 20", width=30)
    createButton([OwnerReports], 5, 0, 2, "Sold Food Items", lambda:generateSoldFoodItemsGraph(), "Calibri 20", width=30)
    createButton([OwnerReports], 6, 0, 2, "Item Trends", lambda:generateItemTrendGraph(), "Calibri 20", width=30)

#############################################################################################################################
#############################################################################################################################
#############################################################################################################################

def floatToPrice(f: float):
    return f"£{format(f, ',.2f')}"

def ordinal(n: int):
    if 11 <= (n % 100) <= 13:
        suffix = 'th'
    else:
        suffix = ['th', 'st', 'nd', 'rd', 'th'][min(n % 10, 4)]
    return str(n) + suffix

def monthToTime(month: int or str, year: int or str):
    if(type(month) == int):
        m = month
    else:
        m = indexFromVal(months, month) + 1
    if(type(year) == int):
        y = year
    else:
        y = int(year)
    return time.mktime(time.struct_time([y, m, 1, 1, 0, 0, 0, 1, 0])) - 1704067200
    
months = ["January", "February", "March", "April", "May", "June", "July", "August", "September", "October", "November", "December"]

def truncateText(text: str, length: int):
    if(len(text) <= length + 1):
        return text
    if(text[length-1]) == ' ':
        return f'{text[0:length-1]}...'
    return f'{text[0:length]}...'

def closedWindow(frame):
    global selectedOrder
    selectedOrder = []
    frame.destroy()

#############################################################################################################################
#############################################################################################################################
#############################################################################################################################