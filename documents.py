from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx import Document
from docx.shared import Inches
from tkinter import messagebox
import time
from time import strftime
import json
import os

from main import cursor
from functions import floatToPrice, ordinal

# Modified from https://github.com/python-openxml/python-docx/issues/433#issuecomment-358566765

def modifyBorder(table, top, bottom, left, right):
    # Get the XML of the table
    tbl = table._tbl
    i = 0
    for cell in tbl.iter_tcs():
        # Get the tcPr, from which we can modify the borders
        tcPr = cell.tcPr
        tcBorders = OxmlElement('w:tcBorders')
        topBorder = OxmlElement('w:top')
        
        # Change the borders
        if(top[i]): topBorder.set(qn('w:val'), 'single')
        else: topBorder.set(qn('w:val'), 'nil')
        
        bottomBorder = OxmlElement('w:bottom')
        if(bottom[i]): bottomBorder.set(qn('w:val'), 'single')
        else: bottomBorder.set(qn('w:val'), 'nil')
        
        leftBorder = OxmlElement('w:left')
        if(left[i]): leftBorder.set(qn('w:val'), 'single')
        else: leftBorder.set(qn('w:val'), 'nil')
        
        rightBorder = OxmlElement('w:right')
        if(right[i]): rightBorder.set(qn('w:val'), 'single')
        else: rightBorder.set(qn('w:val'), 'nil')

        # Apply changes
        tcBorders.append(topBorder)
        tcBorders.append(bottomBorder)
        tcBorders.append(leftBorder)
        tcBorders.append(rightBorder)
        tcPr.append(tcBorders)

        i += 1

# Generate an array of repeated Falses or Trues
def genBrdArr(val: bool, length: int):
    return [val for i in range(0, length)]

# Generate the full border definitions
def genBrds(columns: int, items: int):
    borders = [[], [], [], []]
    borders[0].extend(genBrdArr(False, columns))
    borders[1].extend(genBrdArr(False, columns))
    borders[2].extend(genBrdArr(False, columns))
    borders[3].extend(genBrdArr(False, columns))

    for i in range(0, items):
        borders[0].extend(genBrdArr(True, columns))
        borders[1].extend(genBrdArr(False, columns))
        borders[2].extend(genBrdArr(False, columns))
        borders[3].extend(genBrdArr(False, columns))

        # This whole section is a mess. Making it so that the column value isn't reduced by 1 when there's only 2 columns works. Don't ask me how.
        borders[0].extend(genBrdArr(False, columns - (0 if columns == 2 else 1)))
        borders[1].extend(genBrdArr(False, columns))
        borders[2].extend(genBrdArr(False, columns))
        borders[3].extend(genBrdArr(False, columns))
    return borders

# Generate a word document containing the receipt of an order
def createOrderReceipt(parentTV = None, orderData = {}, order = 0):
    
    from orders import modTypes, sauceDict, picklesDict, appetisers, baos, bentos, bentoSides, classics, classicSides, sides
    from documents import modifyBorder, genBrds
    
    try:   
        # When data is provided that means the receipt is being generated from an order that is currently being placed
        # When data isn't provided, data is coming from existing orders in the database
        if(orderData == {}):
            if(not(parentTV)):
                orderNum = order
            else:
                from toplevels import selectedOrder
                if(not(selectedOrder)):
                    messagebox.showerror("Error", "No order selected", parent=parentTV)
                    return
                orderNum = selectedOrder[0]
            open(f'documents/orderReceipts/{orderNum}.docx')
        else:
            orderNum = 'NA - Temporary receipt'
            raise FileNotFoundError
    except FileNotFoundError:
        doc = Document()
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.1)
            section.bottom_margin = Inches(0.1)
            section.left_margin = Inches(0.1)
            section.right_margin = Inches(0.1)

        doc.add_heading(f'Order {orderNum}', 0)

        doc.add_picture('images/bb_logo_b.png', width=Inches(2))

        if(orderData == {}):
            cursor.execute(f"""SELECT * FROM orders WHERE orderID = {orderNum}""")
            data = cursor.fetchone()
            orderData = json.loads(str(data[2]).removeprefix('b\'').removesuffix('\''))
            orderOverhead = data[5:8]
        else:
            orderOverhead = [int(time.time() - 1704067200), 0, "No note yet"]

        doc.add_heading('Appetisers', level=1)

        items = []

        totalPrice = 0.00

        for d in orderData['appetisers']:
            mods = []

            for i in range(0, len(d['details'][1])):
                modMod = modTypes[appetisers[d['details'][0]]['mod'][i]['modType']][d['details'][1][i]]
                modName = appetisers[d['details'][0]]['mod'][i]['name']
                modFull = f'{modMod} {modName}'
                if(modFull == "Not Well done"): continue
                if(modFull == "Normal Hot"): continue
                if(modName == "Well done"): mods.append(f'{modName}'); continue
                if(modName == "Hot"): mods.append(f'{modMod}'); continue
                if(modFull == "Yes Lemon"): continue
                mods.append(modFull)
            items.append((
                f'{d['count']} {appetisers[d['details'][0]]['name'].removesuffix(' [Veg]').removesuffix(' [Vgn]')}', 
                f'{sauceDict[d['details'][2]].removesuffix(' [Veg]').removesuffix(' [Vgn]')}',
                mods,
                int(d['count']) * appetisers[d['details'][0]]['price'],
                d['note']
            ))

        table = doc.add_table(rows=1, cols=3)
        header = table.rows[0].cells
        header[0].text = 'Qty&Name'
        header[1].text = 'Sauce'
        header[2].text = 'Mod.'
        for qname, sauce, mod, price, note in items:
            totalPrice += price
            row = table.add_row().cells
            row[0].text = qname
            row[1].text = sauce
            for m in mod:
                row[2].text += f'{m}\n'
            row[2].text = row[2].text.removesuffix('\n')
            noteRow = table.add_row().cells
            noteRow[0].merge(noteRow[1])
            noteRow[0].text = note
            noteRow[2].text = floatToPrice(price)
        modifyBorder(table, *genBrds(3, len(items)))


        doc.add_heading('Baos', level=1)
        
        items = []
        for d in orderData['baos']:
            items.append((
                f'{d['count']} {baos[d['details'][0]]['name'].removesuffix(' [Veg]').removesuffix(' [Vgn]')}', 
                f'{f"{modTypes[5][d['details'][2]]} " if d['details'][1] != 0 else ""}{sauceDict[d['details'][1]].removesuffix(' [Veg]').removesuffix(' [Vgn]')}', 
                [f'{modTypes[1][d['details'][3][i]]} {picklesDict[i + 1]}' for i in range(0, 5)],
                int(d['count']) * baos[d['details'][0]]['price'],
                d['note']
            ))

        table = doc.add_table(rows=1, cols=3)
        header = table.rows[0].cells
        header[0].text = 'Qty&Name'
        header[1].text = 'Sauce'
        header[2].text = 'Pickles'
        for qname, sauce, mod, price, note in items:
            totalPrice += price
            row = table.add_row().cells
            row[0].text = qname
            row[1].text = sauce
            for m in mod:
                row[2].text += f'{m}\n'
            row[2].text = row[2].text.removesuffix('\n')
            noteRow = table.add_row().cells
            noteRow[0].merge(noteRow[1])
            noteRow[0].text = note
            noteRow[2].text = floatToPrice(price)
        modifyBorder(table, *genBrds(3, len(items)))
        
        
        doc.add_heading('Bentos', level=1)

        items = []
        for d in orderData['bentos']:
            mods = []
            s1mods = []
            s2mods = []

            for i in range(0, len(d['details'][1])):
                modMod = modTypes[bentos[d['details'][0]]['mod'][i]['modType']][d['details'][1][i]]
                modName = bentos[d['details'][0]]['mod'][i]['name']
                modFull = f'{modMod} {modName}'
                if(modFull == "Normal Hot"): continue
                if(modName == "Hot"): mods.append(f'{modMod}'); continue
                mods.append(modFull)

            for i in range(0, len(d['details'][3])):
                s1modMod = modTypes[bentoSides[d['details'][2]]['mod'][i]['modType']][d['details'][3][i]]
                s1modName = bentoSides[d['details'][2]]['mod'][i]['name']
                s1mods.append(f'{s1modMod} {s1modName}')

            for i in range(0, len(d['details'][5])):
                s2modMod = modTypes[bentoSides[d['details'][4]]['mod'][i]['modType']][d['details'][5][i]]
                s2modName = bentoSides[d['details'][4]]['mod'][i]['name']
                s2mods.append(f'{s2modMod} {s2modName}')

            sauceModFinal = ''
            for i in range(0, len(d['details'][7])):
                sauceModMod = modTypes[sides[d['details'][6]]['mod'][i]['modType']][d['details'][7][i]]  
                sauceModName = sides[d['details'][6]]['mod'][i]['name']

                sauceModFull = f'{sauceModMod} {sauceModName}'
                if(sauceModFull == "Normal Hot"): continue
                if(sauceModFull == "No Onion"): continue
                sauceModFinal += f"{sauceModName}, "
            sauceModFinal = sauceModFinal.removesuffix(', ')

            items.append((
                f'{d['count']} {bentos[d['details'][0]]['name'].removesuffix(' [Veg]').removesuffix(' [Vgn]')}', 
                f'{sauceDict[d['details'][6]].removesuffix(' [Veg]').removesuffix(' [Vgn]')}',
                sauceModFinal,
                mods,
                f'{bentoSides[d['details'][2]]['name'].removesuffix(' [Veg]').removesuffix(' [Vgn]')}',
                s1mods,
                f'{bentoSides[d['details'][4]]['name'].removesuffix(' [Veg]').removesuffix(' [Vgn]')}',
                s2mods,
                int(d['count']) * bentos[d['details'][0]]['price'],
                d['note']
            ))

        table = doc.add_table(rows=1, cols=7)
        header = table.rows[0].cells
        header[0].text = 'Qty&Name'
        header[1].text = 'Sauce'
        header[2].text = 'Mod.'
        header[3].text = 'Side 1'
        header[4].text = 'S1 Mod.'
        header[5].text = 'Side 2'
        header[6].text = 'S2 Mod.'
        for qname, sauce, sauceMod, mod, s1, s1m, s2, s2m, price, note in items:
            totalPrice += price
            row = table.add_row().cells
            row[0].text = qname
            row[1].text = f'{sauce}{f"\n({sauceMod})" if sauceMod else ""}'
            for m in mod:
                row[2].text += f'{m}\n'
            row[2].text = row[2].text.removesuffix('\n')
            row[3].text = s1
            for m1 in s1m:
                row[4].text += f'{m1}\n'
            row[4].text = row[4].text.removesuffix('\n')
            row[5].text = s2
            for m2 in s2m:
                row[6].text += f'{m2}\n'
            row[6].text = row[6].text.removesuffix('\n')
            noteRow = table.add_row().cells
            noteRow[0].merge(noteRow[1])
            noteRow[0].text = note
            noteRow[6].text = floatToPrice(price)
        modifyBorder(table, *genBrds(7, len(items)))

        
        doc.add_heading('Classics', level=1)
        
        items = []
        for d in orderData['classics']:
            mods = []
            sMods = []

            for i in range(0, len(d['details'][1])):
                modMod = modTypes[classics[d['details'][0]]['mod'][i]['modType']][d['details'][1][i]]
                modName = classics[d['details'][0]]['mod'][i]['name']
                modFull = f'{modMod} {modName}'
                if([modMod, modName] == "Normal Hot"): continue
                if(modName == "Hot"): mods.append(f'{modMod}'); continue
                mods.append(modFull)

            for i in range(0, len(d['details'][3])):
                sModMod = modTypes[classicSides[d['details'][2]]['mod'][i]['modType']][d['details'][3][i]]
                sModName = classicSides[d['details'][2]]['mod'][i]['name']
                sMods.append(f'{sModMod} {sModName}')

            items.append((
                f'{d['count']} {classics[d['details'][0]]['name'].removesuffix(' [Veg]').removesuffix(' [Vgn]')}', 
                mods,
                f'{classicSides[d['details'][2]]['name'].removesuffix(' [Veg]').removesuffix(' [Vgn]')}',
                sMods,
                int(d['count']) * classics[d['details'][0]]['price'],
                d['note']
            ))

        table = doc.add_table(rows=1, cols=4)
        header = table.rows[0].cells
        header[0].text = 'Qty&Name'
        header[1].text = 'Mod.'
        header[2].text = 'Side'
        header[3].text = 'S Mod.'
        for qname, mod, s, sm, price, note in items:
            totalPrice += price
            row = table.add_row().cells
            row[0].text = qname
            for m in mod:
                row[1].text += f'{m}\n'
            row[1].text = row[1].text.removesuffix('\n')
            row[2].text = s
            for m in sm:
                row[3].text += f'{m}\n'
            row[3].text = row[3].text.removesuffix('\n')
            noteRow = table.add_row().cells
            noteRow[0].merge(noteRow[1])
            noteRow[0].text = note
            noteRow[3].text = floatToPrice(price)
        modifyBorder(table, *genBrds(4, len(items)))
        
        

        doc.add_heading('Sides', level=1)

        items = []
        for d in orderData['sides']:
            mods = []
            sMods = []

            for i in range(0, len(d['details'][1])):
                modMod = modTypes[sides[d['details'][0]]['mod'][i]['modType']][d['details'][1][i]]
                modName = sides[d['details'][0]]['mod'][i]['name']
                modFull = f'{modMod} {modName}'
                if(modFull == "Normal Hot"): continue
                if(modFull == "Yes Onion"): mods.append(f'Plus {modName}'); continue
                if(modFull == "No Onion"): continue
                mods.append(modFull)

            items.append((
                f'{d['count']} {sides[d['details'][0]]['name'].removesuffix(' [Veg]').removesuffix(' [Vgn]')}', 
                mods,
                int(d['count']) * sides[d['details'][0]]['price'],
                d['note']
            ))

        table = doc.add_table(rows=1, cols=2)
        header = table.rows[0].cells
        header[0].text = 'Qty&Name'
        header[1].text = 'Mod.'
        for qname, mod, price, note in items:
            totalPrice += price
            row = table.add_row().cells
            row[0].text = qname
            for m in mod:
                row[1].text += f'{m}\n'
            row[1].text = row[1].text.removesuffix('\n')
            noteRow = table.add_row().cells
            noteRow[0].text = note
            noteRow[1].text = floatToPrice(price)
        modifyBorder(table, *genBrds(2, len(items)))

        doc.add_paragraph("---\n")

        placementTime = time.localtime(orderOverhead[0] + 1704067200)
        pickupTime = orderOverhead[1]
        doc.add_paragraph(f"Placed: {strftime("%A", placementTime)} {ordinal(int(strftime("%d", placementTime)))} {strftime("%B, %Y at %H:%M", placementTime)}")

        if(pickupTime != 0):
            pickupTime = time.localtime(orderOverhead[1] + 1704067200)
            doc.add_paragraph(f"Pickup Time: {strftime("%A", pickupTime)} {ordinal(int(strftime("%d", pickupTime)))} {strftime("%B, %Y at %H:%M", pickupTime)}")
        else:
            doc.add_paragraph(f"Pickup Time: Not yet specified")

        doc.add_paragraph(f"Note: {orderOverhead[2]}")

        doc.add_paragraph(f"Total price: {floatToPrice(totalPrice)}")

        doc.save(f'documents/orderReceipts/{orderNum}.docx')

    finally:
        if(order == 0):
            if os.system(f'start documents/orderReceipts/"{orderNum}".docx') != 0:
                messagebox.showerror("Error", "Cannot open receipt, file already open!", parent=parentTV)